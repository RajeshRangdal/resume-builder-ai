import os
import subprocess
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import scrolledtext
from mistralai import Mistral

# Function to create the base resume in DOCX format
def create_resume():
    print("Creating resume...")
    document = Document()
    
    # Title (Centered Name and Contact Info)
    name_paragraph = document.add_paragraph()
    name_paragraph.add_run('Enter Name').bold = True
    name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact_info = document.add_paragraph(
        'Phone: +91-XXXXXXXXXX | Email: your-email@example.com | LinkedIn: https://linkedin.com/in/enter-your-profile'
    )
    
    # Summary Section
    document.add_heading('Summary', level=1)
    document.add_paragraph('Add a brief professional summary or personal statement.')
    
    # Education Section
    document.add_heading('Education', level=1)
    document.add_heading('College', level=2)
    document.add_paragraph('Add details about your degree, university, and year of graduation.')
    document.add_heading('School', level=2)
    document.add_paragraph('Add details about your high school, board, and year of passing.')
    
    # Work Experience Section
    document.add_heading('Work Experience', level=1)
    document.add_paragraph('Work experiences will be added based on user input.')
    
    # Projects Section
    document.add_heading('Projects', level=1)
    document.add_paragraph('Projects will be added based on user input.')
    
    # Skills Section
    document.add_heading('Skills', level=1)
    document.add_paragraph('Add a list of your programming languages, cloud platforms, DevOps tools, databases, etc.')
    
    # Certifications Section
    document.add_heading('Certifications', level=1)
    document.add_paragraph('Add details of certifications like AWS, Python, etc.')
    
    # Additional Activities Section
    document.add_heading('Co-curricular and Extracurricular Activities', level=1)
    document.add_paragraph('Add details about any extracurricular activities.')
    
    document.save('resume.docx')
    print('Resume created successfully: resume.docx')

# Function to get structured work experience
def get_work_experience():
    work_experiences = []
    while True:
        title = input("Enter work experience title (e.g., Software Developer at XYZ Ltd.): ").strip()
        duration = input("Enter duration (e.g., June 2020 - Present): ").strip()
        descriptions = []
        print("Enter work descriptions (one per line). Type 'done' when finished:")
        while True:
            desc = input("- ").strip()
            if desc.lower() == 'done':
                break
            descriptions.append(desc)
        work_experiences.append({
            'title': title,
            'duration': duration,
            'descriptions': descriptions
        })
        another = input("Do you want to add another work experience? [y/n]: ").lower().strip()
        if another != 'y':
            break
    return work_experiences

# Function to get structured projects
def get_projects():
    projects = []
    while True:
        title = input("Enter project title: ").strip()
        descriptions = []
        print("Enter project descriptions (one per line). Type 'done' when finished:")
        while True:
            desc = input("- ").strip()
            if desc.lower() == 'done':
                break
            descriptions.append(desc)
        projects.append({
            'title': title,
            'descriptions': descriptions
        })
        another = input("Do you want to add another project? [y/n]: ").lower().strip()
        if another != 'y':
            break
    return projects

# Function to format work experience for LaTeX
def format_work_experience_latex(work_experiences):
    if not work_experiences:
        return r"\item No work experience provided"
    latex_content = []
    for exp in work_experiences:
        exp_latex = f"\\item \\textbf{{{exp['title']}}} ({exp['duration']})"
        if exp['descriptions']:
            exp_latex += "\\begin{itemize}"
            for desc in exp['descriptions']:
                exp_latex += f"\\item {desc}"
            exp_latex += "\\end{itemize}"
        latex_content.append(exp_latex)
    return "\n".join(latex_content)

# Function to format projects for LaTeX
def format_projects_latex(projects):
    if not projects:
        return r"\item No projects provided"
    latex_content = []
    for proj in projects:
        proj_latex = f"\\item \\textbf{{{proj['title']}}}"
        if proj['descriptions']:
            proj_latex += "\\begin{itemize}"
            for desc in proj['descriptions']:
                proj_latex += f"\\item {desc}"
            proj_latex += "\\end{itemize}"
        latex_content.append(proj_latex)
    return "\n".join(latex_content)

# Function to get user input
def get_user_input():
    name = input("Enter your name: ").strip()
    contact_number = input("Enter contact number: ").strip()
    email = input("Enter email: ").strip()
    linkedin = input("Enter LinkedIn URL: ").strip()
    github = input("Enter GitHub URL: ").strip()
    summary = input("Enter your summary: ").strip()

    education = []
    print("Enter your education details. Type 'done' when finished:")
    while True:
        course = input("Course name (or 'done' to finish): ").strip()
        if course.lower() == 'done':
            break
        institution = input("Institution/University name: ").strip()
        marks = input("Marks obtained (%): ").strip()
        duration = input("Duration (e.g., Aug 2019 - June 2023): ").strip()
        education.append({
            'course': course,
            'institution': institution,
            'marks': marks,
            'duration': duration
        })

    work_experiences = get_work_experience()
    projects = get_projects()

    skills = input("Enter your skills (comma-separated): ").strip()
    certifications = input("Enter your certifications (comma-separated): ").strip()
    activities = input("Enter your activities (comma-separated): ").strip()
    languages = input("Enter languages you know (comma-separated): ").strip()

    return (
        name, contact_number, email, linkedin, github, summary,
        education, work_experiences, projects,
        skills, certifications, activities, languages
    )

# Function to create LaTeX resume
def create_latex_resume(name, contact_number, email, linkedin, github, summary, education, work_experiences, projects, skills, certifications, activities, languages):
    # LaTeX preamble with UTF-8 support
    latex_preamble = r"""
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
"""

    # Format work experiences and projects for LaTeX
    work_exp_latex = format_work_experience_latex(work_experiences)
    projects_latex = format_projects_latex(projects)

    # Format education as LaTeX code
    education_latex = ""
    for edu in education:
        education_latex += f"""\\item \\textbf{{{edu['course']}}}\\\\
{edu['institution']} \\hfill {edu['marks']}\\% \\\\
\\hfill {edu['duration']}
"""

    # Format other sections
    skills_latex = r"\item " + r"\item ".join(skills.split(","))
    certifications_latex = r"\item " + r"\item ".join(certifications.split(","))
    activities_latex = r"\item " + r"\item ".join(activities.split(","))
    languages_latex = r"\item " + r"\item ".join(languages.split(","))

    # Rest of the LaTeX template
    latex_content = r"""\documentclass[a4paper,10pt]{article}
\usepackage[left=0.6in,right=0.6in,top=0.5in,bottom=1in]{geometry}
\usepackage{xcolor}
\usepackage{enumitem}
\usepackage[hidelinks]{hyperref}
\usepackage{titlesec}
""" + latex_preamble + r"""
% Define custom colors
\definecolor{darkblue}{RGB}{26,13,171}
\setlength{\parindent}{0pt}

% Set section formatting
\titleformat{\section}{\large\bfseries\color{darkblue}}{}{0em}{}[\titlerule]

% Set hyperlink color
\hypersetup{colorlinks=true, linkcolor=darkblue, urlcolor=darkblue}

\newenvironment{myitemize}{\begin{itemize}[leftmargin=*,labelsep=0.5em,itemsep=0.5em]}{\end{itemize}}

\begin{document}

\pagestyle{empty}

\begin{center}
  \textbf{\LARGE """ + name + r"""}
\end{center}

\begin{center}
  \textbf{Mobile:} """ + contact_number + r""" \quad
  \textbf{Email:} \href{mailto:""" + email + r"""}{""" + email + r"""} \quad
  \href{""" + linkedin + r"""}{LinkedIn} \quad
  \href{""" + github + r"""}{GitHub}
\end{center}

\section*{Summary}
""" + summary + r"""

\section*{Education}
\begin{myitemize}
  """ + education_latex + r"""
\end{myitemize}

\section*{Work Experience}
\begin{itemize}
  """ + work_exp_latex + r"""
\end{itemize}

\section*{Projects}
\begin{itemize}
  """ + projects_latex + r"""
\end{itemize}

\section*{Skills}
\begin{itemize}
  """ + skills_latex + r"""
\end{itemize}

\section*{Courses \& Certifications}
\begin{itemize}
  """ + certifications_latex + r"""
\end{itemize}

\section*{Co-curricular and Extracurricular Activities}
\begin{itemize}
  """ + activities_latex + r"""
\end{itemize}

\section*{Languages}
\begin{itemize}
  """ + languages_latex + r"""
\end{itemize}

\end{document}
"""

    return latex_content

# Function to compile LaTeX to PDF
def compile_latex_to_pdf(latex_content, output_filename='resume.pdf'):
    try:
        with open('resume.tex', 'w', encoding='utf-8') as f:
            f.write(latex_content)
        subprocess.run(['pdflatex', '-interaction=nonstopmode', 'resume.tex'], check=True)
        os.rename('resume.pdf', output_filename)
        print(f'PDF created successfully: {output_filename}')
    except subprocess.CalledProcessError as e:
        print(f'Error in compiling LaTeX to PDF: {e}')
        print('Please ensure you have LaTeX and pdflatex installed on your system.')

# Function to tailor resume based on job description
def tailor_resume():
    def process_tailoring():
        job_description = text_area.get("1.0", tk.END).strip()

        try:
            with open(r"C:\Users\rajes\Resume_Builder\resume.tex", "r", encoding="utf-8") as file:
                resume_content = file.read()
            print("Resume file read successfully.")
        except Exception as e:
            print(f"Error reading resume file: {e}")
            result_label.config(text="Error reading resume file. Check the file path and permissions.")
            return

        summary_match = re.search(r'\\section\*{Summary}(.*?)\\section', resume_content, re.DOTALL)
        skills_match = re.search(r'\\section\*{Skills}(.*?)\\section', resume_content, re.DOTALL)

        if summary_match and skills_match:
            current_summary = summary_match.group(1).strip()
            current_skills = skills_match.group(1).strip()

            api_key = os.environ.get("MISTRAL_API_KEY")
            if not api_key:
                result_label.config(text="MISTRAL_API_KEY not found in environment variables.")
                return

            model = "mistral-tiny"
            client = Mistral(api_key=api_key)

            prompt = f"""
            Job Description: {job_description[:3000]}

            Current Resume Summary: {current_summary[:500]}

            Current Resume Skills: {current_skills[:1000]}

            Please provide a short and professional tailored summary (2-3 sentences) and skills section for this resume based on the job description. Format your response as follows:

            Tailored Summary:
            [Your tailored summary here]

            Tailored Skills:
            [Your tailored skills here]
            """

            try:
                chat_response = client.chat.complete(
                    model=model,
                    messages=[{"role": "user", "content": prompt}]
                )

                ai_response = chat_response.choices[0].message.content
                print("AI Response:", ai_response)

                tailored_summary = re.search(r'Tailored Summary:(.*?)Tailored Skills:', ai_response, re.DOTALL)
                tailored_skills = re.search(r'Tailored Skills:(.*)', ai_response, re.DOTALL)

                if tailored_summary and tailored_skills:
                    tailored_summary = tailored_summary.group(1).strip()
                    tailored_skills = tailored_skills.group(1).strip()

                    skills_list = tailored_skills.split('-')
                    formatted_skills = '\\begin{itemize}\n' + '\n'.join([f'\\item {skill.strip()}' for skill in skills_list if skill.strip()]) + '\n\\end{itemize}'

                    resume_content = resume_content.replace(current_summary, tailored_summary)
                    resume_content = resume_content.replace(current_skills, formatted_skills)

                    with open(r"C:\Users\rajes\Resume_Builder\resume.tex", "w", encoding="utf-8") as file:
                        file.write(resume_content)

                    output_dir = r"C:\Users\rajes\Resume_Builder"
                    input_file = r"C:\Users\rajes\Resume_Builder\resume.tex"
                    try:
                        subprocess.check_output(["pdflatex", f"-output-directory={output_dir}", input_file])
                        result_label.config(text="Resume tailored and saved successfully!")
                    except subprocess.CalledProcessError as e:
                        print(f"Error compiling LaTeX file: {e}")
                        result_label.config(text="Error compiling LaTeX file. Please check the file for errors.")
                else:
                    result_label.config(text="Error in parsing AI response. Please try again.")
            except Exception as e:
                print(f"Error in API request: {e}")
                result_label.config(text="Error in API request. Please check your internet connection and API key.")
        else:
            result_label.config(text="Couldn't find Summary or Skills section in the resume.")

    window = tk.Tk()
    window.title("Resume Tailor")

    text_area = scrolledtext.ScrolledText(window, width=60, height=10)
    text_area.pack(pady=10)

    tailor_button = tk.Button(window, text="Tailor Resume", command=process_tailoring)
    tailor_button.pack(pady=5)

    result_label = tk.Label(window, text="")
    result_label.pack(pady=5)

    window.mainloop()

# Main execution flow
if __name__ == '__main__':
    while True:
        print("\n1. Create Resume (DOCX format)")
        print("2. Create LaTeX Resume")
        print("3. Tailor Resume based on Job Description")
        print("4. Exit")
        choice = input("Enter your choice (1/2/3/4): ")

        if choice == '1':
            create_resume()
        elif choice == '2':
            (name, contact_number, email, linkedin, github, summary,
            education, work_experiences, projects, skills, certifications,
            activities, languages) = get_user_input()

            latex_content = create_latex_resume(name, contact_number, email, linkedin, github, summary,
                                education, work_experiences, projects, skills, certifications,
                                activities, languages)
            compile_latex_to_pdf(latex_content)
        elif choice == '3':
            resume_path = r'C:\Users\rajes\Resume_Builder\resume.tex'
            if not os.path.exists(resume_path):
                print(f"resume.tex not found at {resume_path}. Please create a LaTeX resume first using option 2.")
            else:
                print("Launching Resume Tailoring interface...")
                tailor_resume()
        elif choice == '4':
            print("Exiting the program. Goodbye!")
            break
        else:
            print("Invalid choice. Please select a valid option.")