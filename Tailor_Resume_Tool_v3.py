import os
import subprocess
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tkinter as tk
from tkinter import scrolledtext
from mistralai import Mistral

# Function to get structured work experience
def get_work_experience():
    work_experiences = []
    while True:
        title = input("Enter work experience title (e.g., Software Developer at XYZ Ltd.): ").strip()
        duration = input("Enter duration (e.g., June 2020 - Present): ").strip()
        descriptions = []
        print("Enter work descriptions (one per line). Type 'done' when finished:")
        while True:
            desc = input("- ").strip()
            if desc.lower() == 'done':
                break
            descriptions.append(desc)
        work_experiences.append({
            'title': title,
            'duration': duration,
            'descriptions': descriptions
        })
        another = input("Do you want to add another work experience? [y/n]: ").lower().strip()
        if another != 'y':
            break
    return work_experiences

# Function to get structured projects
def get_projects():
    projects = []
    while True:
        title = input("Enter project title: ").strip()
        descriptions = []
        print("Enter project descriptions (one per line). Type 'done' when finished:")
        while True:
            desc = input("- ").strip()
            if desc.lower() == 'done':
                break
            descriptions.append(desc)
        projects.append({
            'title': title,
            'descriptions': descriptions
        })
        another = input("Do you want to add another project? [y/n]: ").lower().strip()
        if another != 'y':
            break
    return projects

# Function to format work experience for LaTeX
def format_work_experience_latex(work_experiences):
    if not work_experiences:
        return r"\item No work experience provided"
    latex_content = []
    for exp in work_experiences:
        exp_latex = f"\\item \\textbf{{{exp['title']}}} ({exp['duration']})"
        if exp['descriptions']:
            exp_latex += "\\begin{itemize}"
            for desc in exp['descriptions']:
                exp_latex += f"\\item {desc}"
            exp_latex += "\\end{itemize}"
        latex_content.append(exp_latex)
    return "\n".join(latex_content)

# Function to format projects for LaTeX
def format_projects_latex(projects):
    if not projects:
        return r"\item No projects provided"
    latex_content = []
    for proj in projects:
        proj_latex = f"\\item \\textbf{{{proj['title']}}}"
        if proj['descriptions']:
            proj_latex += "\\begin{itemize}"
            for desc in proj['descriptions']:
                proj_latex += f"\\item {desc}"
            proj_latex += "\\end{itemize}"
        latex_content.append(proj_latex)
    return "\n".join(latex_content)

# Function to get user input
def get_user_input():
    name = input("Enter your name: ").strip()
    contact_number = input("Enter contact number: ").strip()
    email = input("Enter email: ").strip()
    linkedin = input("Enter LinkedIn URL: ").strip()
    github = input("Enter GitHub URL: ").strip()
    summary = input("Enter your summary: ").strip()

    education = []
    print("Enter your education details. Type 'done' when finished:")
    while True:
        course = input("Course name (or 'done' to finish): ").strip()
        if course.lower() == 'done':
            break
        institution = input("Institution/University name: ").strip()
        marks = input("Marks obtained (%): ").strip()
        duration = input("Duration (e.g., Aug 2019 - June 2023): ").strip()
        education.append({
            'course': course,
            'institution': institution,
            'marks': marks,
            'duration': duration
        })

    work_experiences = get_work_experience()
    projects = get_projects()

    technical_skills = input("Enter your technical skills (comma-separated): ").strip()
    soft_skills = input("Enter your soft skills (comma-separated): ").strip()
    certifications = input("Enter your certifications (comma-separated): ").strip()
    activities = input("Enter your activities (comma-separated): ").strip()
    languages = input("Enter languages you know (comma-separated): ").strip()

    return (
        name, contact_number, email, linkedin, github, summary,
        education, work_experiences, projects,
        technical_skills, soft_skills, certifications, activities, languages
    )

# Function to create LaTeX resume
def create_latex_resume(name, contact_number, email, linkedin, github, summary, education, work_experiences, projects, technical_skills, soft_skills, certifications, activities, languages):
    # LaTeX preamble with UTF-8 support
    latex_preamble = r"""
\usepackage[utf8]{inputenc}
\usepackage[T1]{fontenc}
"""

    # Format work experiences and projects for LaTeX
    work_exp_latex = format_work_experience_latex(work_experiences)
    projects_latex = format_projects_latex(projects)

    # Format education as LaTeX code
    education_latex = ""
    for edu in education:
        education_latex += f"""\\item \\textbf{{{edu['course']}}}\\\\
{edu['institution']} \\hfill {edu['marks']}\\% \\\\
\\hfill {edu['duration']}
"""

    # Format skills
    technical_skills_latex = r"\begin{itemize}" + "\n" + '\n'.join([f"  \\item {skill.strip()}" for skill in technical_skills.split(",") if skill.strip()]) + "\n\\end{itemize}"
    soft_skills_latex = r"\begin{itemize}" + "\n" + '\n'.join([f"  \\item {skill.strip()}" for skill in soft_skills.split(",") if skill.strip()]) + "\n\\end{itemize}"

    # Format other sections
    certifications_latex = r"\item " + r"\item ".join(certifications.split(","))
    activities_latex = r"\item " + r"\item ".join(activities.split(","))
    languages_latex = r"\item " + r"\item ".join(languages.split(","))

    # Rest of the LaTeX template
    latex_content = r"""\documentclass[a4paper,10pt]{article}
\usepackage[left=0.6in,right=0.6in,top=0.5in,bottom=1in]{geometry}
\usepackage{xcolor}
\usepackage{enumitem}
\usepackage[hidelinks]{hyperref}
\usepackage{titlesec}
""" + latex_preamble + r"""
% Define custom colors
\definecolor{darkblue}{RGB}{26,13,171}
\setlength{\parindent}{0pt}

% Set section formatting
\titleformat{\section}{\large\bfseries\color{darkblue}}{}{0em}{}[\titlerule]

% Set hyperlink color
\hypersetup{colorlinks=true, linkcolor=darkblue, urlcolor=darkblue}

\newenvironment{myitemize}{\begin{itemize}[leftmargin=*,labelsep=0.5em,itemsep=0.5em]}{\end{itemize}}

\begin{document}

\pagestyle{empty}

\begin{center}
  \textbf{\LARGE """ + name + r"""}
\end{center}

\begin{center}
  \textbf{Mobile:} """ + contact_number + r""" \quad
  \textbf{Email:} \href{mailto:""" + email + r"""}{""" + email + r"""} \quad
  \href{""" + linkedin + r"""}{LinkedIn} \quad
  \href{""" + github + r"""}{GitHub}
\end{center}

\section*{Summary}
""" + summary + r"""

\section*{Education}
\begin{myitemize}
  """ + education_latex + r"""
\end{myitemize}

\section*{Work Experience}
\begin{itemize}
  """ + work_exp_latex + r"""
\end{itemize}

\section*{Projects}
\begin{itemize}
  """ + projects_latex + r"""
\end{itemize}

\section*{Skills}
\subsection*{Technical Skills}
""" + technical_skills_latex + r"""

\subsection*{Soft Skills}
""" + soft_skills_latex + r"""

\section*{Courses \& Certifications}
\begin{itemize}
  """ + certifications_latex + r"""
\end{itemize}

\section*{Co-curricular and Extracurricular Activities}
\begin{itemize}
  """ + activities_latex + r"""
\end{itemize}

\section*{Languages}
\begin{itemize}
  """ + languages_latex + r"""
\end{itemize}

\end{document}
"""

    return latex_content

# Function to compile LaTeX to PDF
def compile_latex_to_pdf(latex_content, output_filename='resume.pdf'):
    try:
        with open('resume.tex', 'w', encoding='utf-8') as f:
            f.write(latex_content)
        subprocess.run(['pdflatex', '-interaction=nonstopmode', 'resume.tex'], check=True)
        os.rename('resume.pdf', output_filename)
        print(f'PDF created successfully: {output_filename}')
    except subprocess.CalledProcessError as e:
        print(f'Error in compiling LaTeX to PDF: {e}')
        print('Please ensure you have LaTeX and pdflatex installed on your system.')

def latex_to_docx(latex_content, output_filename='resume.docx'):
    try:
        # Create a new Document
        doc = Document()
        
        # Extract sections using regex
        sections = re.split(r'\\section\*{(.*?)}', latex_content)
        
        # Add name (assuming it's in the first part before any section)
        name_match = re.search(r'\\LARGE\s+(.*?)}', sections[0])
        if name_match:
            name = name_match.group(1)
            heading = doc.add_heading(name, 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add contact information
        contact_info = re.search(r'\\begin{center}(.*?)\\end{center}', sections[0], re.DOTALL)
        if contact_info:
            contact_text = contact_info.group(1)
            # Clean up LaTeX formatting
            contact_text = re.sub(r'\\textbf{|}\\href{.*?}{|}\\quad', '', contact_text)
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            contact_para.add_run(contact_text.strip())
        
        # Process each section
        for i in range(1, len(sections), 2):
            if i + 1 < len(sections):
                section_title = sections[i]
                section_content = sections[i + 1]
                
                # Add section heading
                doc.add_heading(section_title, 1)
                
                # Clean up LaTeX formatting
                content = section_content.strip()
                content = re.sub(r'\\begin{.*?}|\\end{.*?}', '', content)
                content = re.sub(r'\\item\s*', '• ', content)
                content = re.sub(r'\\textbf{(.*?)}', r'\1', content)
                content = re.sub(r'\\subsection\*{(.*?)}', r'\1:', content)
                
                # Add content
                doc.add_paragraph(content)
        
        # Save the document
        doc.save(output_filename)
        print(f'DOCX created successfully: {output_filename}')
        return True
    except Exception as e:
        print(f'Error in converting LaTeX to DOCX: {e}')
        return False

def tailor_resume():
    def process_tailoring():
        output_dir = r"C:\Users\rajes\Resume_Builder"
        input_file = os.path.join(output_dir, "resume.tex")

        job_description = text_area.get("1.0", tk.END).strip()

        try:
            with open(input_file, "r", encoding="utf-8") as file:
                resume_content = file.read()
            print("Resume file read successfully.")
        except Exception as e:
            print(f"Error reading resume file: {e}")
            result_label.config(text="Error reading resume file. Check the file path and permissions.")
            return

        summary_match = re.search(r'\\section\*{Summary}(.*?)\\section', resume_content, re.DOTALL)
        skills_match = re.search(r'\\section\*{Skills}(.*?)\\section', resume_content, re.DOTALL)
        technical_skills_match = re.search(r'\\subsection\*{Technical Skills}(.*?)\\subsection', resume_content, re.DOTALL)
        soft_skills_match = re.search(r'\\subsection\*{Soft Skills}(.*?)\\section', resume_content, re.DOTALL)

        if summary_match and skills_match and technical_skills_match and soft_skills_match:
            current_summary = summary_match.group(1).strip()
            current_technical_skills = technical_skills_match.group(1).strip()
            current_soft_skills = soft_skills_match.group(1).strip()

            api_key = os.environ.get("MISTRAL_API_KEY")
            if not api_key:
                result_label.config(text="MISTRAL_API_KEY not found in environment variables.")
                return

            model = "mistral-tiny"
            client = Mistral(api_key=api_key)

            prompt = f"""
Job Description: {job_description[:3000]}

Current Resume Summary: {current_summary[:500]}

Current Resume Technical Skills: {current_technical_skills[:1000]}

Current Resume Soft Skills: {current_soft_skills[:1000]}

Please provide a short and professional tailored summary (2-3 sentences), a tailored technical skills list, and a tailored soft skills list for this resume based on the job description. Format your response as follows:

Tailored Summary:
[Your tailored summary here]

Tailored Technical Skills:
- Skill 1
- Skill 2
- Skill 3
...

Tailored Soft Skills:
- Skill 1
- Skill 2
- Skill 3
...
"""

            try:
                chat_response = client.chat.complete(
                    model=model,
                    messages=[{"role": "user", "content": prompt}]
                )

                ai_response = chat_response.choices[0].message.content
                print("AI Response:", ai_response)

                tailored_summary = re.search(r'Tailored Summary:(.*?)Tailored Technical Skills:', ai_response, re.DOTALL)
                tailored_technical_skills = re.search(r'Tailored Technical Skills:(.*?)Tailored Soft Skills:', ai_response, re.DOTALL)
                tailored_soft_skills = re.search(r'Tailored Soft Skills:(.*)', ai_response, re.DOTALL)

                if tailored_summary and tailored_technical_skills and tailored_soft_skills:
                    tailored_summary = tailored_summary.group(1).strip()
                    tailored_technical_skills = tailored_technical_skills.group(1).strip()
                    tailored_soft_skills = tailored_soft_skills.group(1).strip()

                    # Format Technical Skills
                    tech_skills_list = re.findall(r'-\s*(.*)', tailored_technical_skills)
                    formatted_technical_skills = '\\begin{itemize}\n' + '\n'.join([f'  \\item {skill.strip()}' for skill in tech_skills_list if skill.strip()]) + '\n\\end{itemize}'

                    # Format Soft Skills
                    soft_skills_list = re.findall(r'-\s*(.*)', tailored_soft_skills)
                    formatted_soft_skills = '\\begin{itemize}\n' + '\n'.join([f'  \\item {skill.strip()}' for skill in soft_skills_list if skill.strip()]) + '\n\\end{itemize}'

                    # Replace Summary
                    resume_content = resume_content.replace(current_summary, tailored_summary)

                    # Replace Technical Skills
                    resume_content = resume_content.replace(current_technical_skills, formatted_technical_skills)

                    # Replace Soft Skills
                    resume_content = resume_content.replace(current_soft_skills, formatted_soft_skills)

                    with open(input_file, "w", encoding="utf-8") as file:
                        file.write(resume_content)

                    try:
                        # Compile PDF
                        subprocess.check_output(["pdflatex", f"-output-directory={output_dir}", input_file])
                        
                        # Generate DOCX
                        docx_output = os.path.join(output_dir, "resume.docx")
                        if latex_to_docx(resume_content, docx_output):
                            result_label.config(text="Resume tailored and saved successfully in both PDF and DOCX formats!")
                        else:
                            result_label.config(text="PDF created successfully, but DOCX conversion failed.")
                    except subprocess.CalledProcessError as e:
                        print(f"Error compiling LaTeX file: {e}")
                        result_label.config(text="Error compiling LaTeX file. Please check the file for errors.")
                else:
                    result_label.config(text="Error in parsing AI response. Please try again.")
            except Exception as e:
                print(f"Error in API request: {e}")
                result_label.config(text="Error in API request. Please check your internet connection and API key.")
        else:
            result_label.config(text="Couldn't find Summary or Skills sections in the resume.")

    window = tk.Tk()
    window.title("Resume Tailor")

    tk.Label(window, text="Enter Job Description:").pack(pady=5)
    text_area = scrolledtext.ScrolledText(window, width=60, height=15)
    text_area.pack(pady=10)

    tailor_button = tk.Button(window, text="Tailor Resume", command=process_tailoring)
    tailor_button.pack(pady=5)

    result_label = tk.Label(window, text="")
    result_label.pack(pady=5)

    window.mainloop()

# Main execution flow
if __name__ == '__main__':
    while True:
        print("\n1. Create LaTeX Resume")
        print("2. Tailor Resume based on Job Description")
        print("3. Exit")
        choice = input("Enter your choice (1/2/3): ")

        if choice == '1':
            (name, contact_number, email, linkedin, github, summary,
            education, work_experiences, projects, technical_skills, soft_skills,
            certifications, activities, languages) = get_user_input()

            latex_content = create_latex_resume(name, contact_number, email, linkedin, github, summary,
                                education, work_experiences, projects, technical_skills, soft_skills,
                                certifications, activities, languages)
            compile_latex_to_pdf(latex_content)
        elif choice == '2':
            resume_path = r'C:\Users\rajes\Resume_Builder\resume.tex'
            if not os.path.exists(resume_path):
                print(f"resume.tex not found at {resume_path}. Please create a LaTeX resume first using option 2.")
            else:
                print("Launching Resume Tailoring interface...")
                tailor_resume()
        elif choice == '3':
            print("Exiting the program. Goodbye!")
            break
        else:
            print("Invalid choice. Please select a valid option.")
